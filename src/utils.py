import PyPDF2
import docx
import os
import logging
from typing import List, Optional
from pathlib import Path

from src.error_handler import (
    DocumentProcessingError, 
    handle_errors, 
    validate_file_type, 
    sanitize_filename,
    safe_execute
)

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document types and extract text content"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.txt']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
    
    @handle_errors(default_response=[])
    def process_document(self, file_path: str) -> List[str]:
        """
        Process a document and return text chunks
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of text chunks
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        # Validate file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise DocumentProcessingError(
                f"File too large: {file_size / 1024 / 1024:.1f}MB (max: {self.max_file_size / 1024 / 1024}MB)",
                filename=os.path.basename(file_path)
            )
        
        # Validate file type
        filename = os.path.basename(file_path)
        if not validate_file_type(filename, self.supported_extensions):
            raise DocumentProcessingError(
                f"Unsupported file type. Supported: {', '.join(self.supported_extensions)}",
                filename=filename
            )
        
        logger.info(f"Processing document: {filename}")
        
        # Extract text based on file type
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
            else:
                raise DocumentProcessingError(
                    f"Unsupported file extension: {file_extension}",
                    filename=filename
                )
            
            if not text or not text.strip():
                raise DocumentProcessingError(
                    "No readable text content found in document",
                    filename=filename
                )
            
            # Split into chunks
            chunks = self._create_chunks(text)
            
            if not chunks:
                raise DocumentProcessingError(
                    "Failed to create text chunks from document",
                    filename=filename
                )
            
            logger.info(f"Successfully processed {filename}: {len(chunks)} chunks created")
            return chunks
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to process document: {str(e)}",
                filename=filename
            )
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise DocumentProcessingError(
                        "PDF is encrypted and cannot be processed",
                        filename=os.path.basename(file_path)
                    )
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"PDF processing error: {str(e)}",
                filename=os.path.basename(file_path)
            )
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            raise DocumentProcessingError(
                f"DOCX processing error: {str(e)}",
                filename=os.path.basename(file_path)
            )
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError(
                "Unable to decode text file with supported encodings",
                filename=os.path.basename(file_path)
            )
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Text file processing error: {str(e)}",
                filename=os.path.basename(file_path)
            )
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if not text or len(text.strip()) < 50:  # Skip very short texts
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            # Calculate end position
            end = min(start + self.chunk_size, text_length)
            
            # Try to break at sentence or paragraph boundary
            chunk_text = text[start:end]
            
            # If not at end of document, try to find better break point
            if end < text_length:
                # Look for sentence endings within last 200 characters
                sentence_endings = ['. ', '! ', '? ', '\n\n']
                best_break = -1
                
                for i in range(len(chunk_text) - 1, max(len(chunk_text) - 200, 0), -1):
                    for ending in sentence_endings:
                        if chunk_text[i:i+len(ending)] == ending:
                            best_break = i + len(ending)
                            break
                    if best_break != -1:
                        break
                
                if best_break != -1:
                    chunk_text = chunk_text[:best_break]
                    end = start + best_break
            
            # Clean and add chunk
            chunk_text = chunk_text.strip()
            if chunk_text and len(chunk_text) > 50:  # Only add substantial chunks
                chunks.append(chunk_text)
            
            # Move to next chunk with overlap
            start = max(end - self.chunk_overlap, start + 1)
            
            # Prevent infinite loop
            if start >= text_length:
                break
        
        return chunks
    
    def get_file_info(self, file_path: str) -> dict:
        """Get basic information about a file"""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            file_stat = os.stat(file_path)
            filename = os.path.basename(file_path)
            
            return {
                "filename": sanitize_filename(filename),
                "size_bytes": file_stat.st_size,
                "size_mb": round(file_stat.st_size / 1024 / 1024, 2),
                "extension": Path(file_path).suffix.lower(),
                "supported": validate_file_type(filename, self.supported_extensions)
            }
        except Exception as e:
            return {"error": f"Failed to get file info: {str(e)}"}
    
    def validate_batch_upload(self, file_paths: List[str]) -> dict:
        """Validate multiple files for batch processing"""
        results = {
            "valid_files": [],
            "invalid_files": [],
            "total_size": 0,
            "warnings": []
        }
        
        for file_path in file_paths:
            file_info = self.get_file_info(file_path)
            
            if "error" in file_info:
                results["invalid_files"].append({
                    "path": file_path,
                    "reason": file_info["error"]
                })
                continue
            
            if not file_info["supported"]:
                results["invalid_files"].append({
                    "path": file_path,
                    "reason": f"Unsupported file type: {file_info['extension']}"
                })
                continue
            
            if file_info["size_bytes"] > self.max_file_size:
                results["invalid_files"].append({
                    "path": file_path,
                    "reason": f"File too large: {file_info['size_mb']}MB"
                })
                continue
            
            results["valid_files"].append(file_info)
            results["total_size"] += file_info["size_bytes"]
        
        # Add warnings
        if results["total_size"] > 100 * 1024 * 1024:  # 100MB total
            results["warnings"].append("Large total file size may take longer to process")
        
        if len(results["valid_files"]) > 20:
            results["warnings"].append("Processing many files may take significant time")
        
        return results